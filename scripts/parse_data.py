import os
import sys
import json
import re
import pandas as pd
import openpyxl
from docx import Document

# Paths
script_dir = os.path.dirname(os.path.abspath(__file__))
workspace_dir = os.path.dirname(script_dir)
sample_dir = os.path.join(workspace_dir, "sample data")
data_out_dir = os.path.join(workspace_dir, "src/data")
os.makedirs(data_out_dir, exist_ok=True)

def find_file_case_insensitive(filename):
    if not os.path.exists(sample_dir):
        return os.path.join(sample_dir, filename)
    target_lower = filename.lower()
    files = os.listdir(sample_dir)
    
    # 1. Try exact match (case-insensitive)
    for f in files:
        if f.lower() == target_lower:
            return os.path.join(sample_dir, f)
            
    # 2. Try match for robot guide spreadsheet (The Newbie Guide / The Simplified WR Guide)
    # Give priority to 'newbie' as the new filename
    if "newbie" in target_lower or "simplified" in target_lower or "new players" in target_lower:
        for f in files:
            f_lower = f.lower()
            if "newbie" in f_lower and f_lower.endswith(".xlsx"):
                print(f"Fuzzy matched robot guide spreadsheet (newbie): {f}")
                return os.path.join(sample_dir, f)
        for f in files:
            f_lower = f.lower()
            if ("simplified" in f_lower or "new players" in f_lower) and f_lower.endswith(".xlsx"):
                print(f"Fuzzy matched robot guide spreadsheet (simplified): {f}")
                return os.path.join(sample_dir, f)

    # 3. Try fuzzy match for tier docx (Tier List Rationales.docx / Tier List Rationales & Disclaimers.docx / Tiers.docx)
    if ("tier" in target_lower or "rationale" in target_lower or "disclaimer" in target_lower) and target_lower.endswith(".docx"):
        for f in files:
            f_lower = f.lower()
            if ("tier" in f_lower or "rationale" in f_lower or "disclaimer" in f_lower) and f_lower.endswith(".docx"):
                print(f"Fuzzy matched tier docx: {f}")
                return os.path.join(sample_dir, f)

    # 4. Try fuzzy match for pilot docx
    if "pilot" in target_lower and target_lower.endswith(".docx"):
        for f in files:
            f_lower = f.lower()
            if "pilot" in f_lower and f_lower.endswith(".docx"):
                print(f"Fuzzy matched pilot docx: {f}")
                return os.path.join(sample_dir, f)

    # 5. Try fuzzy match for specialization docx
    if "specialization" in target_lower and target_lower.endswith(".docx"):
        for f in files:
            f_lower = f.lower()
            if "specialization" in f_lower and f_lower.endswith(".docx"):
                print(f"Fuzzy matched specialization docx: {f}")
                return os.path.join(sample_dir, f)

    # 6. Try fuzzy match for weapon dps xlsx
    if ("dps" in target_lower or "weapon" in target_lower) and target_lower.endswith(".xlsx"):
        for f in files:
            f_lower = f.lower()
            if "dps" in f_lower and f_lower.endswith(".xlsx"):
                print(f"Fuzzy matched DPS spreadsheet: {f}")
                return os.path.join(sample_dir, f)

    # 7. Try fuzzy match for WR tier lists xlsx
    if ("wr tier" in target_lower or "tier list" in target_lower or "tiers" in target_lower) and target_lower.endswith(".xlsx"):
        for f in files:
            f_lower = f.lower()
            if "tier" in f_lower and f_lower.endswith(".xlsx") and "newbie" not in f_lower and "simplified" not in f_lower:
                print(f"Fuzzy matched WR tier lists spreadsheet: {f}")
                return os.path.join(sample_dir, f)

    # 8. Generic word matching fallback
    target_stem = os.path.splitext(target_lower)[0]
    target_ext = os.path.splitext(target_lower)[1]
    target_words = [w for w in re.split(r'\W+', target_stem) if len(w) > 2]
    if target_words:
        for f in files:
            f_lower = f.lower()
            if target_ext and not f_lower.endswith(target_ext):
                continue
            if all(word in f_lower for word in target_words):
                print(f"Fuzzy matched file by words: {f}")
                return os.path.join(sample_dir, f)

    return os.path.join(sample_dir, filename)



# Graceful check for GitHub runner where secrets are not yet configured
if not os.path.exists(sample_dir) or not os.listdir(sample_dir):
    print("Warning: No raw spreadsheets or Word documents found in 'sample data/' directory.")
    required_files = ["tiers.json", "robot_guide.json", "weapons_dps.json", "specializations.json", "pilots.json"]
    all_exist = all(os.path.exists(os.path.join(data_out_dir, f)) for f in required_files)
    if all_exist:
        print("Pre-compiled JSON database files exist in 'src/data/'. Skipping parsing and using existing JSON database.")
        sys.exit(0)
    else:
        print("Error: No raw documents found in 'sample data/' and pre-compiled JSON database is missing in 'src/data/'. Cannot build.", file=sys.stderr)
        sys.exit(1)

# ----------------------------------------------------
# 1. PARSE pilot skill guide.docx
# ----------------------------------------------------
def parse_pilots():
    filepath = find_file_case_insensitive("pilot skill guide.docx")
    doc = Document(filepath)
    
    pilots_data = {
        "intro": "",
        "robots": {
            "Must Use": [],
            "Usually Use": [],
            "Sometimes Use": [],
            "Don't Use": []
        },
        "titans": {
            "Must Use": [],
            "Usually Use": [],
            "Sometimes Use": [],
            "Don't Use": []
        }
    }
    
    current_section = "intro" # intro, robots, titans
    current_category = None # Must Use, Usually Use, Sometimes Use, Don't Use
    
    intro_paragraphs = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        style = p.style.name
        
        # Check section headers
        if style.startswith('Heading') or text.lower() in ["robot pilots", "titan pilots", "intro"]:
            text_clean = text.replace('\u200b', '').strip().lower()
            if "robot pilots" in text_clean:
                current_section = "robots"
                current_category = None
            elif "titan pilots" in text_clean:
                current_section = "titans"
                current_category = None
            elif "intro" in text_clean:
                current_section = "intro"
            continue
            
        # Check categories within sections
        text_clean = text.lower().replace(" ", "")
        matched_cat = None
        if "mustuse" in text_clean:
            matched_cat = "Must Use"
        elif "usuallyuse" in text_clean:
            matched_cat = "Usually Use"
        elif "sometimesuse" in text_clean:
            matched_cat = "Sometimes Use"
        elif "don'tuse" in text_clean or "dontuse" in text_clean:
            matched_cat = "Don't Use"
            
        if matched_cat:
            if current_section == "intro":
                current_section = "robots"
            current_category = matched_cat
            continue
                
        # Parse skills
        if current_section == "intro":
            intro_paragraphs.append(text)
        elif current_section in ["robots", "titans"] and current_category:
            if ":" in text:
                parts = text.split(":", 1)
                skill_name = parts[0].strip()
                desc = parts[1].strip()
                pilots_data[current_section][current_category].append({
                    "name": skill_name,
                    "description": desc
                })
            else:
                # continuation or general comments inside a category
                if pilots_data[current_section][current_category]:
                    pilots_data[current_section][current_category][-1]["description"] += " " + text
                else:
                    pilots_data[current_section][current_category].append({
                        "name": "General Info",
                        "description": text
                    })
                    
    pilots_data["intro"] = "\n".join(intro_paragraphs)
    
    with open(os.path.join(data_out_dir, "pilots.json"), "w", encoding="utf-8") as f:
        json.dump(pilots_data, f, indent=2)
    print("Parsed pilot skill guide.")

# ----------------------------------------------------
# 2. PARSE Specializations guide.docx
# ----------------------------------------------------
def parse_specializations():
    filepath = find_file_case_insensitive("Specializations guide.docx")
    doc = Document(filepath)
    
    specs_data = {
        "intro": "",
        "sections": [],
        "conclusion": ""
    }
    
    intro_paragraphs = []
    conclusion_paragraphs = []
    
    current_section = None
    current_slot = None
    
    # Predefined section headings we expect
    section_headings = [
        "Damage Dealer/Raider (Robot)",
        "Brawler/Tank (Robot)",
        "Support (Robot)",
        "Saboteur (Robot)",
        "Attack/Defense (Robot)",
        "Damage Dealer (Titan)",
        "Brawler/Tank (Titan)",
        "Attack (Titan)",
        "Defense (Titan)",
        "Ultimate (Titan)"
    ]
    
    mode = "skip" # skip, intro, body, conclusion
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        style = p.style.name
        
        # Check headings
        if style.startswith('Heading') or text in section_headings or text in ["Intro", "What specialization class does my robot use?", "Passive upgrades", "Conclusion"]:
            if text in section_headings:
                mode = "body"
                current_section = {
                    "title": text,
                    "description": "",
                    "slots": []
                }
                specs_data["sections"].append(current_section)
                current_slot = None
            elif text == "Conclusion":
                mode = "conclusion"
            elif text == "Intro":
                mode = "intro"
                intro_paragraphs.append(f"### {text}")
            elif text in ["What specialization class does my robot use?", "Passive upgrades"]:
                # These are subheadings in intro
                mode = "intro"
                intro_paragraphs.append(f"\n### {text}")
            continue
            
        # Parse content based on mode
        if mode == "intro":
            intro_paragraphs.append(text)
        elif mode == "conclusion":
            conclusion_paragraphs.append(text)
        elif mode == "body" and current_section:
            if text.lower() in ["first module slot", "second module slot"]:
                current_slot = {
                    "name": text,
                    "content": ""
                }
                current_section["slots"].append(current_slot)
            else:
                if current_slot:
                    current_slot["content"] = (current_slot["content"] + "\n" + text).strip()
                else:
                    current_section["description"] = (current_section["description"] + "\n" + text).strip()
                    
    specs_data["intro"] = "\n".join(intro_paragraphs)
    specs_data["conclusion"] = "\n".join(conclusion_paragraphs)
    
    with open(os.path.join(data_out_dir, "specializations.json"), "w", encoding="utf-8") as f:
        json.dump(specs_data, f, indent=2)
    print("Parsed specializations guide.")

# ----------------------------------------------------
# 3. PARSE Adazahi r_WarRobotsGuide weapon DPS spreadsheet.xlsx
# ----------------------------------------------------
def parse_weapons_dps():
    filepath = find_file_case_insensitive("Adazahi r_WarRobotsGuide weapon DPS spreadsheet.xlsx")
    xl = pd.ExcelFile(filepath)
    
    weapons_data = {}
    
    # We ignore 'Intro' sheet
    sheets = [s for s in xl.sheet_names if s != 'Intro']
    
    for sheet in sheets:
        df = pd.read_excel(filepath, sheet_name=sheet)
        # Columns: Weapon, Burst DPS, Cycle DPS, Range, Notes
        # Clean and replace nan
        df = df.dropna(subset=['Weapon'])
        df = df.fillna("")
        
        weapons_list = []
        for _, row in df.iterrows():
            weapons_list.append({
                "name": str(row["Weapon"]).strip(),
                "burst_dps": row["Burst DPS"],
                "cycle_dps": row["Cycle DPS"],
                "range": str(row["Range"]).strip(),
                "notes": str(row["Notes"]).strip()
            })
            
        weapons_data[sheet] = weapons_list
        
    with open(os.path.join(data_out_dir, "weapons_dps.json"), "w", encoding="utf-8") as f:
        json.dump(weapons_data, f, indent=2)
    print("Parsed weapon DPS spreadsheet.")

# ----------------------------------------------------
# 4. PARSE Tiers.docx and WR tier lists.xlsx
# ----------------------------------------------------
def parse_tiers():
    docx_path = find_file_case_insensitive("Tier List Rationales.docx")
    if not os.path.exists(docx_path):
        docx_path = find_file_case_insensitive("Tier List Rationales & Disclaimers.docx")
    if not os.path.exists(docx_path):
        docx_path = find_file_case_insensitive("Tiers.docx")
        
    if not os.path.exists(docx_path) and os.path.exists(sample_dir):
        for f in os.listdir(sample_dir):
            if f.endswith(".docx") and ("tier" in f.lower() or "rationale" in f.lower() or "disclaimer" in f.lower()):
                docx_path = os.path.join(sample_dir, f)
                break

    if not os.path.exists(docx_path):
        print(f"Warning: Tier list docx not found at '{docx_path}'.")
        if os.path.exists(os.path.join(data_out_dir, "tiers.json")):
            print("Using existing pre-compiled src/data/tiers.json.")
            return
        else:
            raise FileNotFoundError(f"Package not found at '{docx_path}'")

    xlsx_path = find_file_case_insensitive("WR tier lists.xlsx")
    if not os.path.exists(xlsx_path) and os.path.exists(sample_dir):
        for f in os.listdir(sample_dir):
            if f.endswith(".xlsx") and "tier" in f.lower() and "newbie" not in f.lower() and "simplified" not in f.lower():
                xlsx_path = os.path.join(sample_dir, f)
                break

    # Collect all items in Excel that end with '*'
    asterisk_items = set()
    try:
        wb_temp = openpyxl.load_workbook(xlsx_path, data_only=True)
        for sname in wb_temp.sheetnames:
            sheet_temp = wb_temp[sname]
            for r in range(1, sheet_temp.max_row + 1):
                col1_val = sheet_temp.cell(row=r, column=2).value
                if col1_val:
                    items_list = [i.strip() for i in str(col1_val).split(",")]
                    for item in items_list:
                        if item.endswith("*"):
                            asterisk_items.add(item.rstrip("*").strip().lower())
    except Exception as e:
        print(f"Warning: Failed to parse asterisks from WR tier lists.xlsx: {e}")

    # Parse standard tiers and descriptions from docx
    doc = Document(docx_path)
    
    categories = ["Robots", "Titans", "Drones", "Motherships", "Mothership Turrets", "Robot Weapons", "Titan Weapons"]
    tiers_list = ["X", "S", "A", "B", "C", "D", "E", "F", "Z"]
    
    # Normalize category names in docx
    category_mapping = {
        "robots": "Robots",
        "titans": "Titans",
        "drones": "Drones",
        "mothership": "Motherships",
        "motherships": "Motherships",
        "mothership turrets": "Mothership Turrets",
        "robot weapons": "Robot Weapons",
        "titan weapons": "Titan Weapons"
    }
    
    tiers_data = {cat: {t: {"casual_name": "", "items": []} for t in tiers_list} for cat in categories}
    disclaimers = []
    
    current_category = None
    current_tier = None
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        style = p.style.name
        text_lower = text.lower()
        
        # Check category heading
        if text_lower in category_mapping:
            current_category = category_mapping[text_lower]
            current_tier = None
            continue
            
        # Save disclaimers before first category
        if current_category is None:
            if text in ["Tab 1", "Intro"]:
                continue
            disclaimers.append(text)
            continue
            
        # Check tier heading (e.g. "X tier", "X Tier", "S tier")
        if text_lower.endswith(" tier") and text_lower.split(" ")[0].upper() in tiers_list:
            current_tier = text_lower.split(" ")[0].upper()
            continue
            
        # If we have category and tier, parse items
        if current_category and current_tier:
            if ":" in text:
                parts = text.split(":", 1)
                item_name = parts[0].strip()
                desc = parts[1].strip()
                
                # Double check if item name is too long or contains sentences (fail-safe)
                if len(item_name) < 300:
                    clean_item_name = item_name.lower().strip()
                    sub_names = [sn.strip().lower() for sn in clean_item_name.split(",")]
                    has_ast = any(sn in asterisk_items for sn in sub_names)
                    if has_ast and not item_name.endswith("*"):
                        item_name = item_name + "*"
                        
                    tiers_data[current_category][current_tier]["items"].append({
                        "name": item_name,
                        "description": desc
                    })
                else:
                    # Treat as description continuation
                    if tiers_data[current_category][current_tier]["items"]:
                        tiers_data[current_category][current_tier]["items"][-1]["description"] += " " + text
            else:
                # Continuation description
                if tiers_data[current_category][current_tier]["items"]:
                    tiers_data[current_category][current_tier]["items"][-1]["description"] += " " + text
                else:
                    # General notes under a tier
                    tiers_data[current_category][current_tier]["items"].append({
                        "name": "Note",
                        "description": text
                    })

    # Now parse casual tier names from WR tier lists.xlsx
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet = wb.active
    
    # Normalize category names in Excel
    xlsx_category_mapping = {
        "robots": "Robots",
        "titans": "Titans",
        "titan weapons": "Titan Weapons",
        "drones": "Drones",
        "motherships": "Motherships",
        "weapons": "Robot Weapons",
        "mothership turrets": "Mothership Turrets"
    }
    
    excel_current_category = None
    tier_index = 0
    
    for r in range(1, sheet.max_row + 1):
        col0_val = sheet.cell(row=r, column=1).value
        col1_val = sheet.cell(row=r, column=2).value
        
        if col0_val is None:
            continue
            
        col0_str = str(col0_val).strip()
        col0_lower = col0_str.lower()
        
        # Check if it is a category header
        if col0_lower in xlsx_category_mapping:
            excel_current_category = xlsx_category_mapping[col0_lower]
            tier_index = 0
            continue
            
        # If we have a category, match rows to tiers
        if excel_current_category:
            if tier_index < len(tiers_list):
                tier_letter = tiers_list[tier_index]
                casual_name = col0_str
                # Save casual tier name
                tiers_data[excel_current_category][tier_letter]["casual_name"] = casual_name
                tier_index += 1
                
    # If any casual names are missing, use default standard tier letter
    for cat in categories:
        for t in tiers_list:
            if not tiers_data[cat][t]["casual_name"]:
                tiers_data[cat][t]["casual_name"] = f"{t} Tier"
                
    tiers_data["disclaimers"] = disclaimers

    with open(os.path.join(data_out_dir, "tiers.json"), "w", encoding="utf-8") as f:
        json.dump(tiers_data, f, indent=2)
    print("Parsed tier lists and descriptions.")

# ----------------------------------------------------
# 5. PARSE The Simplified WR Guide for New Players.xlsx
# ----------------------------------------------------
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))

def rgb_to_hex(rgb):
    return '#{:02x}{:02x}{:02x}'.format(*rgb)

def interpolate_color(color1, color2, factor):
    rgb1 = hex_to_rgb(color1)
    rgb2 = hex_to_rgb(color2)
    rgb_interp = tuple(int(round(a + factor * (b - a))) for a, b in zip(rgb1, rgb2))
    return rgb_to_hex(rgb_interp)

def extract_rating_colors(wb):
    # Default fallback colors (bright primary/secondary palette)
    rating_colors = {
        "min": "#ff0000",
        "mid": "#ffff00",
        "max": "#00ff00",
        "super": "#00ffff"
    }
    try:
        sheet = None
        for sname in ["Tier 4 Robots", "Tier 4", "Tier 3 Robots", "Tier 3", "Ultimate Robots", "All Titans", "Titans", "Value Rating Criterion"]:
            if sname in wb.sheetnames:
                sheet = wb[sname]
                break
        if sheet is not None:
            cfs = sheet.conditional_formatting
            for cf in cfs:
                for rule in cf.rules:
                    if rule.type == "colorScale" and rule.colorScale:
                        colors = rule.colorScale.color
                        if len(colors) >= 3:
                            c_min = colors[0].rgb
                            c_mid = colors[1].rgb
                            c_max = colors[2].rgb
                            if isinstance(c_min, str) and len(c_min) >= 6:
                                rating_colors["min"] = "#" + c_min[-6:].lower()
                            if isinstance(c_mid, str) and len(c_mid) >= 6:
                                rating_colors["mid"] = "#" + c_mid[-6:].lower()
                            if isinstance(c_max, str) and len(c_max) >= 6:
                                rating_colors["max"] = "#" + c_max[-6:].lower()
                    elif rule.type == "cellIs" and rule.operator in ["greaterThanOrEqual", "equal"] and rule.dxfId is not None:
                        dxf_list = wb._differential_styles
                        if hasattr(dxf_list, "styles") and int(rule.dxfId) < len(dxf_list.styles):
                            style = dxf_list.styles[int(rule.dxfId)]
                            if style.fill and style.fill.fill_type == "solid" and style.fill.start_color:
                                c_super = style.fill.start_color.rgb
                                if isinstance(c_super, str) and len(c_super) >= 6:
                                    rating_colors["super"] = "#" + c_super[-6:].lower()
    except Exception as e:
        print(f"Warning: Failed to extract rating colors from spreadsheet: {e}")
        
    c_min = rating_colors["min"]
    c_mid = rating_colors["mid"]
    c_max = rating_colors["max"]
    c_super = rating_colors["super"]
    
    c_neg1 = interpolate_color(c_min, c_mid, 0.5)
    c_pos1 = interpolate_color(c_mid, c_max, 0.5)
    
    return {
        "<= -2": c_min,
        "-1": c_neg1,
        "0": c_mid,
        "+1": c_pos1,
        "+2": c_max,
        ">= +3": c_super
    }

def parse_robot_guide():
    filepath = find_file_case_insensitive("The Newbie Guide.xlsx")
    if not os.path.exists(filepath):
        filepath = find_file_case_insensitive("The Simplified WR Guide for New Players.xlsx")
    if not os.path.exists(filepath):
        filepath = find_file_case_insensitive("The Simplified WR Guide for New Players, 1-10 Ver_.xlsx")
    
    wb = openpyxl.load_workbook(filepath, data_only=True)

    # 5.1 Parse Changelog (skip header row if present)
    changelog = []
    changes_sheet = wb["Changes Log"]
    for r in range(1, changes_sheet.max_row + 1):
        d_val = changes_sheet.cell(row=r, column=1).value
        txt_val = changes_sheet.cell(row=r, column=2).value
        if d_val and txt_val:
            d_str = str(d_val).strip()
            if d_str.lower() in ["date", "timestamp", "time"]:
                continue
            if "00:00:00" in d_str:
                d_str = d_str.split(" ")[0]
            changelog.append({
                "date": d_str,
                "text": str(txt_val).strip()
            })
            
    # 5.2 Parse Mini Build Guides (support both 'Robot Mini Build Guides' and legacy 'Mini Build Guides')
    builds_sheet_name = "Robot Mini Build Guides" if "Robot Mini Build Guides" in wb.sheetnames else "Mini Build Guides"
    builds_sheet = wb[builds_sheet_name]
    builds = []
    ue_weapon_index = {}

    for r in range(2, builds_sheet.max_row + 1):
        bname = builds_sheet.cell(row=r, column=1).value
        bot = builds_sheet.cell(row=r, column=2).value
        
        bname_str = str(bname).strip() if bname else ""
        bot_str = str(bot).strip() if bot else ""

        if bname_str.startswith("*") or bname_str.startswith("UE Index") or "f2p weapons" in bname_str.lower():
            break
            
        if bname and bot:
            is_ultimate = (bname_str.lower() == "ultimate") or bot_str.lower().startswith("ue ")
            builds.append({
                "build_name": bname_str,
                "robot": bot_str,
                "f2p_weapons": str(builds_sheet.cell(row=r, column=3).value or "").strip(),
                "best_weapons": str(builds_sheet.cell(row=r, column=4).value or "").strip(),
                "drone_options": str(builds_sheet.cell(row=r, column=5).value or "").strip(),
                "pilot": str(builds_sheet.cell(row=r, column=6).value or "").strip(),
                "specialization": str(builds_sheet.cell(row=r, column=7).value or "").strip(),
                "explanation": str(builds_sheet.cell(row=r, column=8).value or "").strip(),
                "is_ultimate": is_ultimate
            })

    # Parse UE Weapon Index (rows 78-83 in Robot Mini Build Guides or wherever quoted index appears)
    for r in range(1, builds_sheet.max_row + 1):
        c1_val = str(builds_sheet.cell(row=r, column=1).value or "").strip()
        c2_val = str(builds_sheet.cell(row=r, column=2).value or "").strip()
        if c1_val.startswith('"') and c1_val.endswith('"'):
            cat = c1_val.strip('"')
            ue_weapon_index[cat] = c2_val

    # 5.3 Parse Roles with fill colors
    roles_sheet = wb["Bot Roles"]
    roles_data = {}
    role_headers = []
    for c in range(1, roles_sheet.max_column + 1):
        h = roles_sheet.cell(row=1, column=c).value
        role_headers.append(str(h).strip() if h else "")
    
    for r in range(2, roles_sheet.max_row + 1):
        bot_name = roles_sheet.cell(row=r, column=1).value
        if not bot_name:
            continue
        bot_str = str(bot_name).strip()
        # Skip description text
        if len(bot_str) > 40:
            continue
            
        bot_roles = []
        for c in range(2, len(role_headers) + 1):
            if c - 1 >= len(role_headers) or not role_headers[c-1]:
                continue
            cell = roles_sheet.cell(row=r, column=c)
            val = cell.value
            fill = cell.fill
            
            color_hex = None
            if fill and fill.fill_type == 'solid':
                color = fill.start_color
                if color:
                    color_hex = color.rgb
                    
            role_type = "none"
            if color_hex:
                clean_hex = color_hex[-6:].upper() if isinstance(color_hex, str) else ""
                if clean_hex == "00FF00": # Green
                    role_type = "primary"
                elif clean_hex == "FFFF00": # Yellow
                    role_type = "secondary"
                    
            # If the value is some footnote, we keep it
            footnote = str(val).strip() if val is not None else ""
            
            if role_type != "none" or footnote:
                bot_roles.append({
                    "role": role_headers[c-1],
                    "type": role_type,
                    "footnote": footnote
                })
        roles_data[bot_str.lower()] = bot_roles

    # 5.3.2 Parse UE Bot Roles with fill colors (if present)
    ue_roles_sheet = wb["UE Bot Roles"] if "UE Bot Roles" in wb.sheetnames else None
    if ue_roles_sheet:
        ue_role_headers = []
        for c in range(1, ue_roles_sheet.max_column + 1):
            h = ue_roles_sheet.cell(row=1, column=c).value
            ue_role_headers.append(str(h).strip() if h else "")

        for r in range(2, ue_roles_sheet.max_row + 1):
            bot_name = ue_roles_sheet.cell(row=r, column=1).value
            if not bot_name:
                continue
            bot_str = str(bot_name).strip()
            if len(bot_str) > 40:
                continue

            bot_roles = []
            for c in range(2, len(ue_role_headers) + 1):
                if c - 1 >= len(ue_role_headers) or not ue_role_headers[c-1]:
                    continue
                cell = ue_roles_sheet.cell(row=r, column=c)
                val = cell.value
                fill = cell.fill

                color_hex = None
                if fill and fill.fill_type == 'solid':
                    color = fill.start_color
                    if color:
                        color_hex = color.rgb

                role_type = "none"
                if color_hex:
                    clean_hex = color_hex[-6:].upper() if isinstance(color_hex, str) else ""
                    if clean_hex == "00FF00":
                        role_type = "primary"
                    elif clean_hex == "FFFF00":
                        role_type = "secondary"

                footnote = str(val).strip() if val is not None else ""
                # In UE Bot Roles, Freezo drone note is marked with single asterisk, map to '**' to align with Bot Roles
                if footnote == "*":
                    footnote = "**"

                if role_type != "none" or footnote:
                    bot_roles.append({
                        "role": ue_role_headers[c-1],
                        "type": role_type,
                        "footnote": footnote
                    })
            roles_data[bot_str.lower()] = bot_roles

    # 5.3.5 Parse Titan Roles with fill colors
    titan_roles_sheet_name = "Titan Roles"
    titan_roles_sheet = wb[titan_roles_sheet_name]
    titan_roles_data = {}
    titan_role_headers = []
    for c in range(1, titan_roles_sheet.max_column + 1):
        h = titan_roles_sheet.cell(row=1, column=c).value
        titan_role_headers.append(str(h).strip() if h else "")
    
    for r in range(2, titan_roles_sheet.max_row + 1):
        titan_name = titan_roles_sheet.cell(row=r, column=1).value
        if not titan_name:
            continue
        titan_str = str(titan_name).strip()
        # Skip description text
        if len(titan_str) > 40:
            continue
            
        titan_roles = []
        for c in range(2, len(titan_role_headers) + 1):
            if c - 1 >= len(titan_role_headers) or not titan_role_headers[c-1]:
                continue
            cell = titan_roles_sheet.cell(row=r, column=c)
            val = cell.value
            fill = cell.fill
            
            color_hex = None
            if fill and fill.fill_type == 'solid':
                color = fill.start_color
                if color:
                    color_hex = color.rgb
                    
            role_type = "none"
            if color_hex:
                clean_hex = color_hex[-6:].upper() if isinstance(color_hex, str) else ""
                if clean_hex == "00FF00": # Green
                    role_type = "primary"
                elif clean_hex == "FFFF00": # Yellow
                    role_type = "secondary"
                    
            footnote = str(val).strip() if val is not None else ""
            
            if role_type != "none" or footnote:
                titan_roles.append({
                    "role": titan_role_headers[c-1],
                    "type": role_type,
                    "footnote": footnote
                })
        titan_roles_data[titan_str.lower()] = titan_roles

    # Helper function to parse side-by-side ratings and score details tables
    def parse_scores_table(sheet, sheet_display_name, lookup_roles):
        left_ratings = {}
        right_details = {}

        for r in range(2, sheet.max_row + 1):
            left_bot = sheet.cell(row=r, column=1).value
            left_rating = sheet.cell(row=r, column=2).value
            if left_bot is not None and left_rating is not None:
                try:
                    left_ratings[str(left_bot).strip()] = int(float(left_rating))
                except (ValueError, TypeError):
                    left_ratings[str(left_bot).strip()] = 0

            right_bot = sheet.cell(row=r, column=4).value
            if right_bot is not None:
                right_bot_str = str(right_bot).strip()
                if sheet.cell(row=r, column=10).value is not None or sheet.cell(row=r, column=5).value is not None:
                    right_details[right_bot_str] = {
                        "longevity": sheet.cell(row=r, column=5).value,
                        "lethality": sheet.cell(row=r, column=6).value,
                        "mobility": sheet.cell(row=r, column=7).value,
                        "utility": sheet.cell(row=r, column=8).value,
                        "accessibility": sheet.cell(row=r, column=9).value,
                        "overall": sheet.cell(row=r, column=10).value,
                        "comments": str(sheet.cell(row=r, column=11).value or "").strip()
                    }

        left_keys_mapped = {k.lower().strip(): (k, v) for k, v in left_ratings.items()}
        result_items = []
        unmatched = []

        def to_num(v):
            if v is None:
                return 0.0
            try:
                return float(v)
            except (ValueError, TypeError):
                return 0.0

        for rname, details in right_details.items():
            rname_clean = rname.lower().strip()
            val_rating = 0
            original_name = rname

            if rname_clean in left_keys_mapped:
                original_name, val_rating = left_keys_mapped[rname_clean]
            else:
                unmatched.append(f"item '{rname}' in {sheet_display_name}")

            roles = lookup_roles.get(original_name.lower().strip(), [])

            result_items.append({
                "name": original_name,
                "sheet": sheet_display_name,
                "value_rating": val_rating,
                "scores": {
                    "longevity": to_num(details["longevity"]),
                    "lethality": to_num(details["lethality"]),
                    "mobility": to_num(details["mobility"]),
                    "utility": to_num(details["utility"]),
                    "accessibility": to_num(details["accessibility"]),
                    "overall": to_num(details["overall"])
                },
                "roles": roles,
                "comments": details["comments"]
            })

        return result_items, unmatched

    # 5.4 Parse Tiers/Value ratings and detailed scores
    # Combine regular robots ('Tier 4' + 'Tier 3') and 'Ultimate Robots'
    robots_data = []
    unmatched_rating_rows = []

    robot_sheet_names = []
    if "Tier 4 Robots" in wb.sheetnames:
        robot_sheet_names.extend(["Tier 4 Robots", "Tier 3 Robots"])
    elif "Tier 4" in wb.sheetnames:
        robot_sheet_names.extend(["Tier 4", "Tier 3"])

    if "Ultimate Robots" in wb.sheetnames:
        robot_sheet_names.append("Ultimate Robots")

    for sheet_name in robot_sheet_names:
        sheet = wb[sheet_name]
        items, unmatched = parse_scores_table(sheet, sheet_name, roles_data)
        robots_data.extend(items)
        unmatched_rating_rows.extend(unmatched)

    # 5.5 Parse Titans sheet (support 'All Titans' and 'Titans')
    titan_sheet_name = "All Titans" if "All Titans" in wb.sheetnames else "Titans"
    titans_sheet = wb[titan_sheet_name]
    titans_data, titan_unmatched = parse_scores_table(titans_sheet, titan_sheet_name, titan_roles_data)
    unmatched_rating_rows.extend(titan_unmatched)

    if unmatched_rating_rows:
        joined = "\n  - ".join(unmatched_rating_rows)
        raise ValueError(
            "The robot guide detail table contains names that are missing from "
            f"the value-rating table:\n  - {joined}"
        )
        
    # Parse footnotes from Bot Roles, UE Bot Roles, and Titan Roles sheets
    roles_footnotes = []
    seen_footnotes = set()
    footnote_sheets = [roles_sheet, titan_roles_sheet]
    if ue_roles_sheet:
        footnote_sheets.append(ue_roles_sheet)

    for sheet in footnote_sheets:
        for r in range(1, sheet.max_row + 1):
            for c in range(1, sheet.max_column + 1):
                val = sheet.cell(row=r, column=c).value
                if val:
                    val_str = str(val).strip()
                    if re.match(r'^\*+\s*[A-Za-z]', val_str):
                        # Normalize single asterisk Freezo note if present
                        if "freezo" in val_str.lower() and val_str.startswith("*") and not val_str.startswith("**"):
                            val_str = "*" + val_str
                        if val_str not in seen_footnotes:
                            seen_footnotes.add(val_str)
                            roles_footnotes.append(val_str)
    # Sort footnotes by asterisk count (*, **, ***, ...)
    roles_footnotes.sort(key=lambda x: len(re.match(r'^\*+', x).group(0)) if re.match(r'^\*+', x) else 0)

    rating_colors = extract_rating_colors(wb)

    # Output unified guide JSON
    guide_json = {
        "changelog": changelog,
        "builds": builds,
        "robots": robots_data,
        "titans": titans_data,
        "camelot": [],
        "footnotes": roles_footnotes,
        "ue_weapon_index": ue_weapon_index,
        "rating_colors": rating_colors
    }
    
    with open(os.path.join(data_out_dir, "robot_guide.json"), "w", encoding="utf-8") as f:
        json.dump(guide_json, f, indent=2)
    print("Parsed robot guide spreadsheet.")


# ----------------------------------------------------
# MAIN EXECUTION
# ----------------------------------------------------
if __name__ == "__main__":
    print("Starting processing of Google Drive / sample data content...")
    parse_pilots()
    parse_specializations()
    parse_weapons_dps()
    parse_tiers()
    parse_robot_guide()
    print("All parsing completed successfully. Output JSONs saved to src/data/")
